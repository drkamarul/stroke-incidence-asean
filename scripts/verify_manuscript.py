"""Verify the manuscript."""
import os
import re
from docx import Document

OUT_DOCX = "/sessions/eloquent-hopeful-franklin/mnt/test_agents_claude_cowork/outputs/stroke_incidence_manuscript_asean.docx"

size = os.path.getsize(OUT_DOCX)
print(f"File size: {size} bytes ({size / 1024:.1f} KB)")
assert size > 200 * 1024

doc = Document(OUT_DOCX)
parts = []
for para in doc.paragraphs:
    parts.append(para.text)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                parts.append(para.text)
all_text = "\n".join(parts)

cite_matches = re.findall(r"\[Cite:[^\]]+\]", all_text)
print(f"[Cite: placeholders remaining: {len(cite_matches)}")
assert len(cite_matches) == 0

citation_re = re.compile(r"\(([^()]*?\d{4}[^()]*?)\)")
raw_citations = citation_re.findall(all_text)


def normalize(cite):
    cite = cite.strip()
    if re.match(r"^\d{4}", cite):
        return None
    if re.search(r"\d{4}\s*[-–]\s*\d{4}", cite):
        return None
    if "million in" in cite or "primary dataset" in cite:
        return None
    m = re.search(r"(.+?),?\s*(\d{4})\s*$", cite)
    if not m:
        return None
    head = m.group(1).strip().rstrip(",")
    year = m.group(2)
    head = re.sub(r"\bet al\.?\b", "", head)
    head = head.replace(".", " ")
    head = re.sub(r"\s+", " ", head).strip().rstrip(",").strip()
    head = re.split(r",| and ", head)[0].strip()
    return f"{head} {year}"


in_text_keys = set()
for raw in raw_citations:
    for piece in raw.split(";"):
        norm = normalize(piece)
        if norm:
            in_text_keys.add(norm)

print(f"\nIn-text citation keys ({len(in_text_keys)}):")
for k in sorted(in_text_keys):
    print(f"  - {k}")

ref_start = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "References":
        ref_start = i
        break

ref_keys = set()
if ref_start is not None:
    for para in doc.paragraphs[ref_start + 1:]:
        t = para.text.strip()
        if not t:
            continue
        m = re.match(r"^(.+?)\s*\((\d{4})\)", t)
        if m:
            head = m.group(1).strip().rstrip(",")
            year = m.group(2)
            head = re.split(r",| and ", head)[0].strip()
            ref_keys.add(f"{head} {year}")

print(f"\nReference list keys ({len(ref_keys)}):")
for k in sorted(ref_keys):
    print(f"  - {k}")

missing = in_text_keys - ref_keys
unused = ref_keys - in_text_keys
print(f"\nIn-text without reference: {sorted(missing)}")
print(f"References not cited in-text: {sorted(unused)}")

# Final pass: also check the cited markdown
md_path = "/sessions/eloquent-hopeful-franklin/mnt/test_agents_claude_cowork/drafts/discussion_draft_cited.md"
with open(md_path) as f:
    md = f.read()
md_cites = re.findall(r"\[Cite:[^\]]+\]", md)
print(f"\n[Cite: placeholders in cited markdown: {len(md_cites)}")
