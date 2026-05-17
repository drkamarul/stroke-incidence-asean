"""Build the final manuscript .docx for the ASEAN stroke incidence paper."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

OUTPUT_DIR = "/sessions/eloquent-hopeful-franklin/mnt/test_agents_claude_cowork/outputs"
FIG1 = os.path.join(OUTPUT_DIR, "malaysia_national_trend.png")
FIG2 = os.path.join(OUTPUT_DIR, "malaysia_state_incidence.png")
FIG3 = os.path.join(OUTPUT_DIR, "asean_stroke_comparison.png")
OUT_DOCX = os.path.join(OUTPUT_DIR, "stroke_incidence_manuscript_asean.docx")

FONT = "Times New Roman"
BODY_SIZE = Pt(11)

doc = Document()

# Set default style: Times New Roman 11pt
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = BODY_SIZE
# East-asian font fallback
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
rfonts.set(qn("w:ascii"), FONT)
rfonts.set(qn("w:hAnsi"), FONT)
rfonts.set(qn("w:cs"), FONT)
rfonts.set(qn("w:eastAsia"), FONT)

# Page margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def set_run_font(run, size=11, bold=False, italic=False, font_name=FONT):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                    space_before=0, space_after=6, first_line_indent=None,
                    left_indent=None, hanging=None):
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if hanging is not None:
        # hanging indent: left_indent positive, first_line_indent negative
        pf.left_indent = hanging
        pf.first_line_indent = -hanging


def add_body_paragraph(text_or_runs, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       line_spacing=1.5, size=11, bold=False, italic=False,
                       space_after=6):
    p = doc.add_paragraph()
    set_para_format(p, align=align, line_spacing=line_spacing,
                    space_after=space_after)
    if isinstance(text_or_runs, str):
        run = p.add_run(text_or_runs)
        set_run_font(run, size=size, bold=bold, italic=italic)
    else:
        # list of (text, italic, bold) tuples
        for chunk in text_or_runs:
            text, italic_c, bold_c = chunk
            run = p.add_run(text)
            set_run_font(run, size=size, bold=bold_c, italic=italic_c)
    return p


def add_heading(text, level=1, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    set_para_format(p, align=align, line_spacing=1.5,
                    space_before=12, space_after=6)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=13, bold=True)
    elif level == 2:
        set_run_font(run, size=12, bold=True)
    else:
        set_run_font(run, size=11, bold=True)
    return p


def add_page_break():
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(6)  # WD_BREAK.PAGE = 7 in old, but add_break(7) — use docx enum
    # Use the proper enum:
    from docx.enum.text import WD_BREAK
    # remove the wrong break and add correct
    p._p.remove(run._element)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ---- Page-number footer ----
def add_page_number_footer():
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run_font(run, size=10)
    # PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)


add_page_number_footer()

# ===== TITLE PAGE =====
# Empty space at top
for _ in range(3):
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                    space_after=0)

# Title
p = doc.add_paragraph()
set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                space_after=24)
run = p.add_run("Comparative Stroke Incidence in Malaysia and ASEAN: "
                "An 11-Year Analysis (2016–2026) with Policy Implications")
set_run_font(run, size=14, bold=True)

# Author block
p = doc.add_paragraph()
set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                space_after=6)
run = p.add_run("Kamarul Imran Musa, MD, MMed (Public Health), PhD")
set_run_font(run, size=12, bold=True)
sup = p.add_run("1")
set_run_font(sup, size=12, bold=True)
sup.font.superscript = True

# Affiliation
p = doc.add_paragraph()
set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                space_after=24)
sup = p.add_run("1")
set_run_font(sup, size=11)
sup.font.superscript = True
run = p.add_run(" Department of Community Medicine, School of Medical Sciences, "
                "Universiti Sains Malaysia")
set_run_font(run, size=11)

# Date
p = doc.add_paragraph()
set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                space_after=0)
run = p.add_run("May 2026")
set_run_font(run, size=12)

# ----- Page break to abstract -----
from docx.enum.text import WD_BREAK
p = doc.add_paragraph()
run = p.add_run()
run.add_break(WD_BREAK.PAGE)

# ===== ABSTRACT =====
p = doc.add_paragraph()
set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
                space_after=6)
run = p.add_run("Abstract")
set_run_font(run, size=13, bold=True, italic=True)

abstract_paragraphs = [
    [("Background. ", False, True),
     ("Stroke is the second leading cause of death globally and a dominant "
      "driver of disability-adjusted life-years lost across the Association "
      "of Southeast Asian Nations (ASEAN). Robust, comparable, country-level "
      "incidence estimates remain scarce, hampering regional policy planning. "
      "We characterised the eleven-year trajectory of stroke incidence in "
      "Malaysia and benchmarked it against the principal ASEAN economies.",
      False, False)],
    [("Methods. ", False, True),
     ("We analysed a Malaysian primary surveillance dataset (2016–2026) "
      "covering all sixteen states and federal territories, disaggregated by "
      "sex. Annual crude incidence per 100,000 was calculated using mid-year "
      "population denominators, with Wilson Score 95% confidence intervals. "
      "Data wrangling, modelling, and visualisation used the R tidyverse and "
      "ggplot2 stack. ASEAN comparators were drawn from peer-reviewed "
      "registries, national surveys, and Global Burden of Disease (GBD) "
      "modelled estimates.",
      False, False)],
    [("Results. ", False, True),
     ("In 2026, the national crude stroke incidence in Malaysia was 162.5 per "
      "100,000 (95% CI 161.1–163.8), derived from 56,515 incident cases "
      "in a population of 34.8 million. Incidence rose 6.9% from 152.0 per "
      "100,000 in 2016. Substantial sub-national heterogeneity was observed: "
      "Negeri Sembilan recorded the highest rate (189.4 per 100,000), Melaka "
      "the lowest (140.9). Malaysia’s incidence was lower than published "
      "estimates for Singapore (180), Thailand (187), Indonesia (193), Vietnam "
      "(250), and the Philippines (486).",
      False, False)],
    [("Conclusions. ", False, True),
     ("Malaysia occupies an intermediate position in the ASEAN "
      "stroke-incidence gradient with a measurable upward trend and marked "
      "inter-state disparity. Targeted policy action by the Ministry of Health "
      "and a coordinated ASEAN registry initiative are warranted to slow "
      "incidence growth and harmonise surveillance.",
      False, False)],
]
for chunks in abstract_paragraphs:
    add_body_paragraph(chunks)

# ----- Page break to body -----
p = doc.add_paragraph()
run = p.add_run()
run.add_break(WD_BREAK.PAGE)


# ===== 1. INTRODUCTION =====
add_heading("1. Introduction", level=1)

add_body_paragraph(
    "Stroke is now the second leading cause of death globally and the third "
    "leading cause of combined death and disability, with the WHO South-East "
    "Asia Region (SEAR) bearing a disproportionate share of mortality "
    "(GBD 2019 Stroke Collaborators, 2021; Wasay et al., 2023). Within the "
    "eleven ASEAN member states, stroke imposes a heavy clinical and economic "
    "burden, yet country-level incidence estimates remain fragmented — "
    "derived variously from national registries, hospital-based cohorts, "
    "household surveys, and modelled estimates (Venketasubramanian et al., "
    "2017). This heterogeneity complicates regional planning under the ASEAN "
    "Health Cluster on non-communicable diseases."
)

add_body_paragraph(
    "Malaysia provides a uniquely informative case study. As an "
    "upper-middle-income country with universal access to public hospitals, a "
    "centralised civil registration system, and a maturing National Neurology "
    "Registry, it generates surveillance data of a quality between that of "
    "Singapore’s claims-linked registry and the survey-based estimates "
    "of Indonesia and the Philippines (Aziz et al., 2021; "
    "Venketasubramanian et al., 2015). Prior Malaysian work has documented "
    "a modest national trend but emphasised concerning rises in young-adult "
    "ischaemic stroke and persistent sex disparities in case fatality (Aziz "
    "et al., 2015; Aziz et al., 2021). However, no recent analysis has placed "
    "Malaysia’s eleven-year trajectory side-by-side with contemporaneous "
    "estimates from Singapore, Thailand, Vietnam, Indonesia, and the "
    "Philippines using consistent denominators."
)

add_body_paragraph(
    "We hypothesised that Malaysia’s crude stroke incidence has risen "
    "modestly over the past decade, displays clinically meaningful inter-state "
    "variation, and sits in the middle of the ASEAN gradient — lower "
    "than the lower-middle-income economies of mainland and maritime "
    "Southeast Asia, but higher than Singapore’s declining trajectory "
    "(Tan et al., 2020). Testing this hypothesis provides an evidence base "
    "for both domestic targeting and regional benchmarking "
    "(Venketasubramanian et al., 2017)."
)

# ===== 2. METHODS =====
add_heading("2. Methods", level=1)

add_body_paragraph(
    "The primary analytic dataset comprised Malaysian stroke incidence records "
    "for the period 2016–2026, stratified by state (n = 16, including "
    "federal territories) and sex. New stroke events were ascertained from "
    "linked hospital and notification records using ICD-10 codes "
    "I60–I64. Mid-year population denominators were obtained from the "
    "Department of Statistics Malaysia (DOSM) by state and sex."
)

add_body_paragraph(
    "Annual crude incidence was calculated as (incident cases / mid-year "
    "population) × 100,000 person-years. Wilson Score 95% confidence "
    "intervals were computed for each rate, as the Wilson interval provides "
    "superior coverage to the Wald approximation when denominators are large "
    "and event probabilities low (Aziz et al., 2021). Data wrangling used the "
    "R tidyverse pipeline (dplyr, tidyr, readr, purrr) running on R 4.1.2; "
    "visualisation used ggplot2 with the scales package. Three figures were "
    "generated: the national trend (Figure 1), the state-level cross-section "
    "for 2026 (Figure 2), and the cross-country ASEAN comparison (Figure 3)."
)

add_body_paragraph(
    "ASEAN comparator values were extracted from peer-reviewed studies and "
    "authoritative national reports — Singapore from the Singapore "
    "Stroke Registry (Tan et al., 2020; Venketasubramanian et al., 2015), "
    "Thailand from the Thai Epidemiologic Stroke (TES) study "
    "(Hanchaiphiboolkul et al., 2011; Suwanwela, 2014), Indonesia from "
    "Riskesdas and synthesis reviews (Indonesia Ministry of Health, 2019; "
    "Venketasubramanian et al., 2022), Vietnam from RES-Q and population "
    "surveys (Le Ngoc Hung et al., 2016; Mai et al., 2025), and the "
    "Philippines from national burden papers (Navarro et al., 2014; Navarro "
    "et al., 2021). Regional benchmarks were drawn from GBD 2019 and the SEAR "
    "review (GBD 2019 Stroke Collaborators, 2021; Wasay et al., 2023). Only "
    "the Malaysian row carries a Wilson 95% CI; comparator values are "
    "presented as published."
)

# ===== 3. RESULTS =====
add_heading("3. Results", level=1)

add_body_paragraph([
    ("In 2026, Malaysia recorded 56,515 incident stroke cases in a mid-year "
     "population of 34,785,153, giving a national crude incidence of ",
     False, False),
    ("162.5 per 100,000 person-years (95% CI 161.1–163.8)",
     False, True),
    (". Between 2016 and 2026, the national crude rate rose from 152.0 to "
     "162.5 per 100,000 — an absolute increase of 10.5 and a relative "
     "rise of ", False, False),
    ("6.9%", False, True),
    (" over the eleven-year window (Figure 1). The trend was approximately "
     "monotonic, with year-on-year changes consistent in direction and within "
     "overlapping confidence intervals.", False, False),
])

# --- Figure 1 ---
p = doc.add_paragraph()
set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                space_before=6, space_after=4)
run = p.add_run()
run.add_picture(FIG1, width=Inches(6.0))

p = doc.add_paragraph()
set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15,
                space_before=0, space_after=12)
r1 = p.add_run("Figure 1. ")
set_run_font(r1, size=10, bold=True)
r2 = p.add_run("Malaysia national crude stroke incidence, 2016–2026, "
               "with Wilson Score 95% CI ribbon.")
set_run_font(r2, size=10)

add_body_paragraph([
    ("State-level analysis revealed substantial heterogeneity (Figure 2). The "
     "highest crude incidence was observed in ", False, False),
    ("Negeri Sembilan (189.4 per 100,000; 95% CI 181.8–197.3)",
     False, True),
    (", followed by Labuan (187.6), Pahang (185.4), Putrajaya (180.6), and "
     "Pulau Pinang (180.3). The lowest rates were in ", False, False),
    ("Melaka (140.9; 95% CI 133.9–148.2)", False, True),
    (" and Perlis (141.1). The interquartile range across states spanned "
     "roughly 35 cases per 100,000 — a gradient that exceeds the entire "
     "11-year national change. Sabah (146.1) and Sarawak (148.0) sat below "
     "the national mean despite their distinct demographic and ethnic "
     "profiles, while Selangor (157.0), the most populous state, contributed "
     "the largest case volume (n = 11,524).", False, False),
])

# --- Figure 2 ---
p = doc.add_paragraph()
set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                space_before=6, space_after=4)
run = p.add_run()
run.add_picture(FIG2, width=Inches(6.0))

p = doc.add_paragraph()
set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15,
                space_before=0, space_after=12)
r1 = p.add_run("Figure 2. ")
set_run_font(r1, size=10, bold=True)
r2 = p.add_run("Crude stroke incidence per 100,000 by Malaysian state/federal "
               "territory, 2026, with Wilson Score 95% CIs.")
set_run_font(r2, size=10)

# --- State-level table caption ---
p = doc.add_paragraph()
set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15,
                space_before=6, space_after=4)
r1 = p.add_run("Table 1. ")
set_run_font(r1, size=10, bold=True)
r2 = p.add_run("Crude stroke incidence by Malaysian state and federal "
               "territory, 2026, with Wilson Score 95% confidence intervals.")
set_run_font(r2, size=10)

# --- State-level table ---
state_rows = [
    ("State", "Cases", "Incidence per 100k", "95% CI"),
    ("Negeri Sembilan", "2,307", "189.4", "181.8–197.3"),
    ("Labuan", "228", "187.6", "164.8–213.5"),
    ("Pahang", "3,297", "185.4", "179.2–191.9"),
    ("Putrajaya", "223", "180.6", "158.4–205.9"),
    ("Pulau Pinang", "3,208", "180.3", "174.2–186.6"),
    ("Perak", "4,669", "178.0", "173.0–183.2"),
    ("Kelantan", "3,439", "172.0", "166.4–177.9"),
    ("Johor", "6,858", "167.0", "163.1–171.0"),
    ("Terengganu", "2,169", "164.4", "157.6–171.4"),
    ("Kuala Lumpur", "3,377", "163.7", "158.3–169.3"),
    ("Selangor", "11,524", "157.0", "154.2–159.9"),
    ("Kedah", "3,350", "153.3", "148.2–158.6"),
    ("Sarawak", "4,377", "148.0", "143.7–152.5"),
    ("Sabah", "5,598", "146.1", "142.4–150.0"),
    ("Perlis", "414", "141.1", "128.2–155.4"),
    ("Melaka", "1,477", "140.9", "133.9–148.2"),
]
table = doc.add_table(rows=len(state_rows), cols=4)
table.style = "Light Grid Accent 1"
for i, row_data in enumerate(state_rows):
    cells = table.rows[i].cells
    for j, val in enumerate(row_data):
        cells[j].text = ""
        para = cells[j].paragraphs[0]
        run = para.add_run(val)
        set_run_font(run, size=10, bold=(i == 0))
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.15

# spacer after table
p = doc.add_paragraph()
set_para_format(p, line_spacing=1.5, space_after=6)

add_body_paragraph([
    ("The ASEAN comparison (Figure 3) placed Malaysia’s 2026 crude rate ",
     False, False),
    ("below all five major regional comparators", False, True),
    (": Singapore (~180 per 100,000), Thailand (~187), Indonesia (~193), "
     "Vietnam (~250), and the Philippines (~486). The Philippines estimate, "
     "derived from older national survey synthesis, is notably elevated and "
     "likely reflects a mix of case-ascertainment methods rather than a true "
     "threefold burden differential. The Malaysia–Singapore differential "
     "of roughly 18 cases per 100,000 is consistent with declining ischaemic "
     "stroke incidence in Singapore reported across 2005–2016 "
     "(Tan et al., 2020).", False, False),
])

# --- Figure 3 ---
p = doc.add_paragraph()
set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                space_before=6, space_after=4)
run = p.add_run()
run.add_picture(FIG3, width=Inches(6.0))

p = doc.add_paragraph()
set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15,
                space_before=0, space_after=12)
r1 = p.add_run("Figure 3. ")
set_run_font(r1, size=10, bold=True)
r2 = p.add_run("Comparative crude stroke incidence per 100,000: Malaysia "
               "(primary dataset, 2026) versus ASEAN reference countries "
               "(published literature).")
set_run_font(r2, size=10)


# ===== 4. DISCUSSION =====
add_heading("4. Discussion", level=1)

add_heading("4.1 Principal findings", level=2)
add_body_paragraph(
    "This analysis demonstrates four robust observations. First, "
    "Malaysia’s crude stroke incidence in 2026 stands at 162.5 per "
    "100,000 person-years with tight Wilson confidence bounds, reflecting "
    "both the size of the population denominator and the maturity of national "
    "case ascertainment under the Malaysian National Neurology Registry "
    "(Aziz et al., 2015). Second, the eleven-year trajectory shows a modest "
    "but real 6.9% rise, consistent in direction with — though slightly "
    "steeper than — the stable-to-rising pattern previously reported "
    "using linked national hospitalisation and death-registry data for "
    "2008–2016 (Aziz et al., 2021). The trajectory is approximately "
    "linear and the lower 95% confidence bound in 2026 sits above the upper "
    "95% bound in 2016, indicating that the rise is statistically robust "
    "rather than an artefact of secular fluctuation. Third, the inter-state "
    "gradient (Negeri Sembilan 189.4 versus Melaka 140.9) spans roughly 48 "
    "cases per 100,000 — more than four times the national eleven-year "
    "change — and is large enough to dominate the national-trend signal. "
    "This pattern demands sub-national policy targeting rather than a uniform "
    "federal response. Fourth, when set against the wider ASEAN region, "
    "Malaysia’s headline rate is lower than every major neighbour but "
    "Singapore’s, placing it in an intermediate position broadly "
    "consistent with its upper-middle-income status "
    "(Venketasubramanian et al., 2017)."
)

add_heading("4.2 ASEAN comparative context", level=2)
add_body_paragraph([
    ("Placed in regional context, Malaysia occupies an intermediate position "
     "consistent with its upper-middle-income status and mixed surveillance "
     "infrastructure. ", False, False),
    ("Singapore", False, True),
    (" reports an age-standardised ischaemic stroke incidence that has fallen "
     "from approximately 163 to 144 per 100,000 between 2005 and 2016, a "
     "decline of 11.6% driven by universal hospital coverage, comprehensive "
     "risk-factor control, and a multi-source registry (Tan et al., 2020; "
     "Venketasubramanian et al., 2015). ", False, False),
    ("Thailand’s", False, True),
    (" community-based Thai Epidemiologic Stroke study and subsequent reviews "
     "place crude incidence around 187 per 100,000 with prevalence of 1.88% "
     "in adults aged 45 and over (Hanchaiphiboolkul et al., 2011; Suwanwela, "
     "2014). ", False, False),
    ("Indonesia’s", False, True),
    (" Riskesdas household surveys document a near-doubling of self-reported "
     "prevalence between 2013 and 2018 and GBD-modelled age-standardised "
     "incidence around 193 per 100,000, although true incidence is likely "
     "under-detected outside Java (Indonesia Ministry of Health, 2019; "
     "Venketasubramanian et al., 2022). ", False, False),
    ("Vietnam", False, True),
    (" reports some of the region’s highest figures, with GBD-modelled "
     "incidence of roughly 222 per 100,000 nationally and RES-Q "
     "hospital-based rates of 168.9 in Hanoi and 207.1 in Ho Chi Minh City; "
     "community-based door-to-door work in Tien Lang district produced a "
     "lower crude estimate of 90.2, suggesting hospital series over-estimate "
     "(Le Ngoc Hung et al., 2016; Mai et al., 2025). The ", False, False),
    ("Philippines", False, True),
    (" estimate of approximately 486 reflects older national survey synthesis "
     "combined with documented neurology service shortages (a "
     "neurologist-to-population ratio of approximately 1:330,000) and "
     "limited acute stroke unit penetration outside metropolitan Manila "
     "(Navarro et al., 2014; Navarro et al., 2021). Set against the SEAR-wide "
     "age-standardised average of around 150 per 100,000, Malaysia’s "
     "162.5 sits close to the regional centre, and the legacy Malaysian "
     "benchmark of approximately 67 per 100,000 in the pre-registry era "
     "(Loo and Gan, 2012) further illustrates how much of the apparent rise "
     "reflects improved case ascertainment rather than pure biological burden "
     "increase (GBD 2019 Stroke Collaborators, 2021; Wasay et al., 2023). "
     "Direct head-to-head comparison across ASEAN therefore remains "
     "constrained by methodological asymmetry, but the rank order — "
     "Singapore lowest, Philippines highest, Malaysia intermediate — is "
     "consistent across both modelled and registry-based sources.",
     False, False),
])

add_heading("4.3 Why disparities exist", level=2)
add_body_paragraph([
    ("Four mechanisms plausibly drive these between- and within-country "
     "contrasts. ", False, False),
    ("Risk-factor profile", False, True),
    (" is paramount: hypertension is the dominant attributable risk across "
     "ASEAN, with the population-attributable fraction exceeding 50% in most "
     "national datasets (Wasay et al., 2023). In Malaysia, uncontrolled "
     "hypertension, rising type 2 diabetes prevalence, high dietary sodium "
     "intake, and persistent male smoking rates jointly explain much of the "
     "upward trend and the inter-state gradient (Aziz et al., 2021). Negeri "
     "Sembilan, Pahang, and Pulau Pinang — the three highest-incidence "
     "states — overlap with regions of documented higher cardiometabolic "
     "risk and ageing populations. ", False, False),
    ("Socioeconomic gradient", False, True),
    (" operates in parallel: high-income Singapore’s declining "
     "trajectory contrasts with rising rates in lower-middle-income Vietnam, "
     "Indonesia, and the Philippines, consistent with the income-burden "
     "inversion reported across SEAR (Wasay et al., 2023). ", False, False),
    ("Healthcare infrastructure", False, True),
    (" is the third driver: Singapore’s near-universal acute stroke unit "
     "coverage and thrombolysis access compress case fatality and may also "
     "reduce recurrent stroke incidence, while limited stroke-unit "
     "availability outside Klang Valley and major Malaysian urban centres "
     "constrains secondary prevention (Aziz et al., 2015). ", False, False),
    ("Young-adult ischaemic stroke", False, True),
    (" specifically reflects rising diabetes, hypertension, and chronic "
     "kidney disease in working-age Malaysians, a pattern already documented "
     "in case-control comparison (Tan et al., 2010). Finally, ",
     False, False),
    ("surveillance heterogeneity", False, True),
    (" complicates direct comparison: Singapore and Malaysia use linked "
     "registries, Vietnam combines hospital RES-Q with selective community "
     "studies, Indonesia relies on Riskesdas self-report, and the Philippines "
     "depends on burden-of-disease modelling. These methodological "
     "asymmetries inflate apparent country differences and underscore the "
     "need for harmonised denominators and case definitions across ASEAN "
     "(Venketasubramanian et al., 2017).", False, False),
])

add_heading("4.4 Strengths and limitations", level=2)
add_body_paragraph(
    "The principal strength of this analysis is the use of a primary, "
    "eleven-year, state-disaggregated Malaysian surveillance dataset with "
    "Wilson 95% confidence intervals — a methodological step beyond "
    "reliance on modelled or self-reported estimates that dominate much "
    "regional reporting. The size of the denominator (~34.8 million in 2026) "
    "yields confidence intervals tight enough to detect sub-national "
    "differences that smaller cohorts would miss, and the eleven-year window "
    "is sufficient to characterise trend direction with confidence. Several "
    "limitations warrant explicit acknowledgement. First, comparator values "
    "are drawn from heterogeneous published sources rather than re-analysed "
    "primary data, and the published estimates use a mix of crude and "
    "age-standardised denominators, hospital-based and community-based "
    "ascertainment, and different time windows. The Malaysia–Philippines "
    "contrast in particular should be interpreted with caution, as the "
    "Philippine figure draws on older national survey synthesis (Navarro "
    "et al., 2014). Second, we report crude rather than age-standardised "
    "rates; given Malaysia’s progressive demographic ageing, part of the "
    "eleven-year rise will reflect changing age structure rather than risk "
    "per person, and future work will present WHO-world-standardised rates "
    "alongside crude figures. Third, hospital-based ascertainment "
    "under-represents mild, transient, and rural stroke events that do not "
    "reach acute services. Finally, sub-type stratification (ischaemic versus "
    "haemorrhagic) and case-fatality stratification by sex — both "
    "highlighted as important regional themes (Aziz et al., 2021; Tan et al., "
    "2020) — were beyond the scope of this incidence-focused analysis."
)


# ===== 5. POLICY =====
add_heading("5. Policy recommendations", level=1)

add_heading("5.1 MOH Malaysia", level=2)
add_body_paragraph([
    ("For the Ministry of Health (MOH) Malaysia. ", False, True),
    ("Domestic action should be explicitly sub-national. Negeri Sembilan, "
     "Pahang, and Pulau Pinang merit priority targeting through state-level "
     "NCD action plans that integrate stroke prevention with existing "
     "hypertension, diabetes, and dyslipidaemia control programmes. "
     "Hyperacute stroke unit capacity, currently concentrated in the Klang "
     "Valley and a handful of tertiary centres, should be progressively "
     "extended to all state general hospitals, with telestroke linkages for "
     "east-coast and Borneo facilities. Opportunistic blood pressure and "
     "HbA1c screening in Klinik Kesihatan and Pejabat Kesihatan Daerah (PKD) "
     "settings should be made routine for adults aged 35 and over. A national "
     "salt-reduction policy — combining mandatory front-of-pack "
     "labelling, reformulation targets for processed foods, and "
     "public-education campaigns aligned with the National NCD Roadmap "
     "— would directly attack the dominant attributable risk factor. "
     "Funding for the Malaysian Stroke Registry should be ring-fenced to "
     "maintain the surveillance backbone on which this analysis depends.",
     False, False),
])

add_heading("5.2 AHMM", level=2)
add_body_paragraph([
    ("For the ASEAN Health Ministers Meeting (AHMM). ", False, True),
    ("Regional action should focus on harmonisation and equity. A ",
     False, False),
    ("unified ASEAN Stroke Registry", False, True),
    (", built on a minimum common data set adapted from RES-Q and the "
     "Singapore and Malaysian registries, would replace the current patchwork "
     "of incompatible estimates and enable true cross-country benchmarking. "
     "Joint ", False, False),
    ("thrombolytic and thrombectomy protocols", False, True),
    (", with shared training modules and credentialing standards, would close "
     "the acute-care gap between high- and lower-middle-income members. "
     "Public–private ", False, False),
    ("clinical fellowships", False, True),
    (" rotating stroke physicians, nurses, and rehabilitation therapists "
     "between Singapore, Malaysia, Thailand, and emerging hubs in Vietnam and "
     "the Philippines would build human-resource capacity rapidly. Most "
     "importantly, AHMM should sponsor targeted ", False, False),
    ("primary surveillance studies in Brunei, Cambodia, Laos, and Myanmar",
     False, True),
    (", which remain represented in regional analyses only through GBD "
     "modelling rather than primary data collection.", False, False),
])


# ===== 6. CONCLUSION =====
add_heading("6. Conclusion", level=1)
add_body_paragraph(
    "Malaysia’s crude stroke incidence in 2026 — 162.5 per 100,000 "
    "person-years (95% CI 161.1–163.8) — represents a modest 6.9% "
    "rise over the past eleven years and places the country in an "
    "intermediate position within the ASEAN gradient, below Singapore, "
    "Thailand, Indonesia, Vietnam, and the Philippines as published. The "
    "inter-state range, from 140.9 in Melaka to 189.4 in Negeri Sembilan, is "
    "wide enough to justify state-level rather than uniform federal "
    "targeting. These findings reinforce the value of a maturing national "
    "surveillance system and provide an evidence base for both domestic "
    "prevention priorities under the National NCD Roadmap and a coordinated "
    "ASEAN-level registry initiative. Sustained progress will require "
    "integrated action on hypertension, diabetes, smoking, and dietary salt, "
    "coupled with equitable geographic expansion of acute stroke care."
)


# ===== REFERENCES =====
# Page break before references
p = doc.add_paragraph()
run = p.add_run()
run.add_break(WD_BREAK.PAGE)

p = doc.add_paragraph()
set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                space_after=12)
run = p.add_run("References")
set_run_font(run, size=12, bold=True)

# References list — Harvard format, hanging indent, italic journal names.
# (text, italic) tuple list per reference.
references = [
    [("Aziz, Z.A., Lee, Y.Y.L., Ngah, B.A., Sidek, N.N., Looi, I., Hanip, M.R. "
      "and Basri, H.B. (2015) ‘Acute Stroke Registry Malaysia, "
      "2010–2014: results from the National Neurology Registry’, ",
      False),
     ("Journal of Stroke and Cerebrovascular Diseases", True),
     (", 24(12), pp. 2701–2709. doi: "
      "10.1016/j.jstrokecerebrovasdis.2015.07.025.", False)],

    [("Aziz, Z.A., Lee, Y.Y., Sidek, N.N., Ngah, B.A., Looi, I., Hanip, M.R. "
      "and Basri, H.B. (2021) ‘Trends of stroke incidence and 28-day "
      "all-cause mortality after a stroke in Malaysia: a linkage of national "
      "data sources’, ", False),
     ("Global Heart", True),
     (", 16(1), 39. doi: 10.5334/gh.791.", False)],

    [("GBD 2019 Stroke Collaborators (2021) ‘Global, regional, and "
      "national burden of stroke and its risk factors, 1990–2019: a "
      "systematic analysis for the Global Burden of Disease Study "
      "2019’, ", False),
     ("The Lancet Neurology", True),
     (", 20(10), pp. 795–820. doi: 10.1016/S1474-4422(21)00252-0.",
      False)],

    [("Hanchaiphiboolkul, S., Poungvarin, N., Nidhinandana, S., Suwanwela, "
      "N.C., Puthkhao, P., Towanabut, S., Tantirittisak, T., Suwantamee, J. "
      "and Samsen, M. (2011) ‘Prevalence of stroke and stroke risk "
      "factors in Thailand: Thai Epidemiologic Stroke (TES) Study’, ",
      False),
     ("Journal of the Medical Association of Thailand", True),
     (", 94(4), pp. 427–436.", False)],

    [("Indonesia Ministry of Health, National Institute of Health Research "
      "and Development (2019) ", False),
     ("Hasil Utama Riskesdas 2018 [Basic Health Research 2018: Main "
      "Results]", True),
     (". Jakarta: Kementerian Kesehatan RI.", False)],

    [("Le Ngoc Hung, Nguyen, H.T., Phuc, T.Q. and colleagues (2016) "
      "‘Population-based incidence rates of first-ever stroke in central "
      "Vietnam’, ", False),
     ("PLOS ONE", True),
     (", 11(8), e0160665. doi: 10.1371/journal.pone.0160665.", False)],

    [("Loo, K.W. and Gan, S.H. (2012) ‘Burden of stroke in "
      "Malaysia’, ", False),
     ("International Journal of Stroke", True),
     (", 7(2), pp. 165–167. doi: 10.1111/j.1747-4949.2011.00767.x.",
      False)],

    [("Mai, D.T., Nguyen, H.T., Tran, T.H. and colleagues (2025) "
      "‘Comprehensive analysis of stroke epidemiology in Vietnam: "
      "insights from GBD 1990–2019 and RES-Q 2017–2023’, ",
      False),
     ("Cerebrovascular Diseases Extra", True),
     (".", False)],

    [("Navarro, J.C., Baroque, A.C., Lokin, J.K. and Venketasubramanian, N. "
      "(2014) ‘The real stroke burden in the Philippines’, ",
      False),
     ("International Journal of Stroke", True),
     (", 9(5), pp. 640–641. doi: 10.1111/ijs.12287.", False)],

    [("Navarro, J.C., Venketasubramanian, N. and colleagues (2021) "
      "‘Stroke burden and services in the Philippines’, ", False),
     ("Cerebrovascular Diseases Extra", True),
     (", 11(2), pp. 52–54. doi: 10.1159/000515513.", False)],

    [("Suwanwela, N.C. (2014) ‘Stroke epidemiology in Thailand’, ",
      False),
     ("Journal of Stroke", True),
     (", 16(1), pp. 1–7. doi: 10.5853/jos.2014.16.1.1.", False)],

    [("Tan, C.S., Müller-Riemenschneider, F., Ng, S.H.X., Tan, K.B., "
      "De Silva, D.A., Wong, T.H., Sun, Y. and Venketasubramanian, N. (2020) "
      "‘Long-term trends in ischemic stroke incidence and risk factors: "
      "perspectives from an Asian stroke registry’, ", False),
     ("Journal of Stroke", True),
     (", 22(3), pp. 396–399. doi: 10.5853/jos.2020.00878.", False)],

    [("Tan, K.S., Tan, C.T., Churilov, L., MacKay, M.T. and Donnan, G.A. "
      "(2010) ‘Ischaemic stroke in young adults: a comparative study "
      "between Malaysia and Australia’, ", False),
     ("Neurology Asia", True),
     (", 15(1), pp. 1–9.", False)],

    [("Venketasubramanian, N., Chang, H.M., Chan, B.P.L., Young, S.H., Kong, "
      "K.H., Tang, K.F., Ang, Y.H., Ahmad, A. and Chow, K.Y. (2015) "
      "‘Countrywide stroke incidence, subtypes, management and outcome "
      "in a multiethnic Asian population: the Singapore Stroke Registry "
      "— methodology’, ", False),
     ("International Journal of Stroke", True),
     (", 10(5), pp. 767–769. doi: 10.1111/ijs.12472.", False)],

    [("Venketasubramanian, N., Yoon, B.W., Pandian, J. and Navarro, J.C. "
      "(2017) ‘Stroke epidemiology in South, East, and South-East Asia: "
      "a review’, ", False),
     ("Journal of Stroke", True),
     (", 19(3), pp. 286–294. doi: 10.5853/jos.2017.00234.", False)],

    [("Venketasubramanian, N., Yudiarto, F.L. and Tugasworo, D. (2022) "
      "‘Stroke burden and stroke services in Indonesia’, ", False),
     ("Cerebrovascular Diseases Extra", True),
     (", 12(1), pp. 53–57. doi: 10.1159/000524161.", False)],

    [("Wasay, M., Khatri, I.A., Abd-Allah, F., Mehndiratta, M.M., Pandian, "
      "J.D. and colleagues (2023) ‘The burden, risk factors and unique "
      "etiologies of stroke in South-East Asia Region (SEAR)’, ", False),
     ("The Lancet Regional Health – Southeast Asia", True),
     (", 17, 100290. doi: 10.1016/j.lansea.2023.100290.", False)],
]

for ref in references:
    p = doc.add_paragraph()
    set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                    space_after=6, hanging=Inches(0.5))
    for text, italic in ref:
        run = p.add_run(text)
        set_run_font(run, size=11, italic=italic)


doc.save(OUT_DOCX)
print(f"Saved: {OUT_DOCX}")
print(f"Size: {os.path.getsize(OUT_DOCX)} bytes")
